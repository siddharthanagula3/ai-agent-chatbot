
try:
    import streamlit as st
    STREAMLIT_AVAILABLE = True
except ImportError:
    print("Streamlit is not available. Falling back to CLI mode.")
    STREAMLIT_AVAILABLE = False

from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationalRetrievalChain
from langchain.vectorstores import Pinecone
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.memory import ConversationBufferMemory
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
import pinecone
import os
import fitz  # PyMuPDF
import docx  # python-docx

# --- API Keys and Initialization ---
openai_api_key = os.getenv("OPENAI_API_KEY")
pinecone_api_key = os.getenv("PINECONE_API_KEY")
pinecone_env = os.getenv("PINECONE_ENV")

pinecone.init(api_key=pinecone_api_key, environment=pinecone_env)
index_name = "langchain-chatbot"

ebd = OpenAIEmbeddings(openai_api_key=openai_api_key)
index = pinecone.Index(index_name)

# --- LLM and Memory ---
chat = ChatOpenAI(temperature=0, openai_api_key=openai_api_key)
memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

# --- Helper: Extract text from PDF, TXT, DOCX ---
def extract_text(file):
    if file.name.endswith(".pdf"):
        doc = fitz.open(stream=file.read(), filetype="pdf")
        return "\n".join(page.get_text() for page in doc)
    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    elif file.name.endswith(".docx"):
        d = docx.Document(file)
        return "\n".join(p.text for p in d.paragraphs if p.text.strip())
    else:
        return ""

# --- Helper: Index text into Pinecone ---
def index_document(text):
    index.delete(delete_all=True, namespace="default")
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    docs = [Document(page_content=chunk) for chunk in splitter.split_text(text)]
    Pinecone.from_documents(docs, embedding=ebd, index_name=index_name, namespace="default")

# --- Retrieval and QA Chain ---
vectorstore = Pinecone(index, ebd.embed_query, "text", namespace="default")
retriever = vectorstore.as_retriever()
qa_chain = ConversationalRetrievalChain.from_llm(llm=chat, retriever=retriever, memory=memory)

if STREAMLIT_AVAILABLE:
    st.set_page_config(page_title="Legal AI Assistant", layout="wide")
    st.title("📚 Legal AI Assistant")

    with st.sidebar:
        st.header("Upload Document")
        uploaded_file = st.file_uploader("Choose a file (PDF, TXT, DOCX)", type=["pdf", "txt", "docx"])
        if uploaded_file:
            with st.spinner("Indexing document..."):
                raw_text = extract_text(uploaded_file)
                index_document(raw_text)
            st.success("Document indexed successfully!")

    st.markdown("Ask me legal questions based on your uploaded document.")

    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    user_query = st.text_input("Ask your question:", key="query")
    if user_query:
        with st.spinner("Generating response..."):
            result = qa_chain.run(user_query)
            st.session_state.chat_history.append((user_query, result))

    if st.session_state.chat_history:
        for q, a in reversed(st.session_state.chat_history):
            st.markdown(f"**You:** {q}")
            st.markdown(f"**AI:** {a}")

    with st.expander("View Full Conversation History"):
        for q, a in st.session_state.chat_history:
            st.markdown(f"**Q:** {q}")
            st.markdown(f"**A:** {a}")
else:
    print("\n📚 Legal AI Assistant - CLI Mode\n")
    chat_history = []
    while True:
        query = input("Ask a legal question (or type 'exit'): ")
        if query.lower() == 'exit':
            break
        answer = qa_chain.run(query)
        chat_history.append((query, answer))
        print(f"AI: {answer}\n")
